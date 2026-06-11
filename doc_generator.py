"""
Document generator using Node.js docx templates.
Falls back to python-docx if Node.js not available.
"""
import json
import os
import subprocess
import tempfile
import sys
from datetime import datetime


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
GEN_DIR  = os.path.join(BASE_DIR, "doc_generators")


def _node_available():
    try:
        r = subprocess.run(["node", "--version"], capture_output=True, timeout=5)
        return r.returncode == 0
    except Exception:
        return False


def _num_to_words(num):
    """Convert number to Russian words."""
    num = round(num)
    ones = ['','один','два','три','четыре','пять','шесть','семь','восемь','девять',
            'десять','одиннадцать','двенадцать','тринадцать','четырнадцать','пятнадцать',
            'шестнадцать','семнадцать','восемнадцать','девятнадцать']
    tens = ['','','двадцать','тридцать','сорок','пятьдесят',
            'шестьдесят','семьдесят','восемьдесят','девяносто']
    hundreds = ['','сто','двести','триста','четыреста','пятьсот',
                'шестьсот','семьсот','восемьсот','девятьсот']
    th_f = ['','одна тысяча','две тысячи','три тысячи','четыре тысячи',
            'пять тысяч','шесть тысяч','семь тысяч','восемь тысяч','девять тысяч']
    if num == 0: return 'ноль рублей 00 копеек'
    result = ''
    th = num // 1000
    rem = num % 1000
    if 0 < th < 10:
        result += th_f[th] + ' '
    elif th >= 10:
        t2 = th // 10; o2 = th % 10
        result += (tens[t2] + ' ' if t2 > 1 else '') + (ones[10+o2] if t2==1 else ones[o2]) + ' тысяч '
    h = rem // 100; t = (rem % 100) // 10; o = rem % 10
    if h: result += hundreds[h] + ' '
    if t == 1: result += ones[10+o] + ' '
    else:
        if t > 1: result += tens[t] + ' '
        if o: result += ones[o] + ' '
    return result.strip() + ' рублей 00 копеек'


def _format_date_ru(date_str):
    """Convert dd.mm.yyyy to 'dd месяц yyyy г.'"""
    months = ['января','февраля','марта','апреля','мая','июня',
              'июля','августа','сентября','октября','ноября','декабря']
    try:
        parts = date_str.split('.')
        d, m, y = int(parts[0]), int(parts[1]), int(parts[2])
        return f"{d} {months[m-1]} {y}"
    except Exception:
        return date_str


def generate_act(data: dict, output_path: str) -> bool:
    """Generate Акт выполненных работ."""
    # Format date for act title
    data = dict(data)
    if 'date' in data:
        data['date'] = _format_date_ru(data['date'])

    if _node_available() and os.path.exists(os.path.join(GEN_DIR, 'make_act.js')):
        return _run_node('make_act.js', data, output_path)
    else:
        return _generate_act_python(data, output_path)


def generate_invoice(data: dict, output_path: str) -> bool:
    """Generate Счёт на оплату."""
    if _node_available() and os.path.exists(os.path.join(GEN_DIR, 'make_invoice.js')):
        return _run_node('make_invoice.js', data, output_path)
    else:
        return _generate_invoice_python(data, output_path)


def generate_upd(data: dict, output_path: str) -> bool:
    """Generate УПД / Счёт-фактура."""
    if _node_available() and os.path.exists(os.path.join(GEN_DIR, 'make_upd.js')):
        return _run_node('make_upd.js', data, output_path)
    else:
        return _generate_act_python(data, output_path)  # fallback


def _run_node(script: str, data: dict, output_path: str) -> bool:
    try:
        tmp_in  = tempfile.mktemp(suffix='.json')
        tmp_out = tempfile.mktemp(suffix='.docx')
        with open(tmp_in, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False)
        script_path = os.path.join(GEN_DIR, script)
        # Patch script to use our temp files
        env = os.environ.copy()
        env['DOC_DATA_FILE'] = tmp_in
        env['DOC_OUT_FILE']  = tmp_out
        # Run with patched paths
        js_code = open(script_path, encoding='utf-8').read()
        js_code = js_code.replace(
            "fs.readFileSync('/tmp/doc_data.json'",
            f"fs.readFileSync('{tmp_in.replace(chr(92), '/')}'"
        ).replace(
            "fs.writeFileSync('/tmp/output_doc.docx'",
            f"fs.writeFileSync('{tmp_out.replace(chr(92), '/')}'"
        )
        tmp_script = tempfile.mktemp(suffix='.js')
        with open(tmp_script, 'w', encoding='utf-8') as f:
            f.write(js_code)
        result = subprocess.run(
            ['node', tmp_script],
            capture_output=True, timeout=30,
            cwd=GEN_DIR
        )
        for f in [tmp_in, tmp_script]:
            try: os.remove(f)
            except: pass
        if result.returncode == 0 and os.path.exists(tmp_out):
            import shutil
            shutil.move(tmp_out, output_path)
            return True
        return False
    except Exception as e:
        print(f"Node error: {e}")
        return False


def _generate_act_python(data: dict, output_path: str) -> bool:
    """Fallback: generate act using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        for s in doc.sections:
            s.top_margin = Cm(2); s.bottom_margin = Cm(2)
            s.left_margin = Cm(2); s.right_margin = Cm(1.5)
        # Simple fallback
        p = doc.add_paragraph()
        r = p.add_run(f"Акт № {data.get('act_number','')} от {data.get('date','')} г.")
        r.bold = True; r.font.size = Pt(14)
        doc.add_paragraph()
        for key, lbl in [('contractor_name','Исполнитель'),('client_name','Заказчик')]:
            p = doc.add_paragraph()
            r1 = p.add_run(f"{lbl}: "); r1.bold = True
            p.add_run(data.get(key,''))
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f"Услуга: {data.get('service_desc','')}")
        doc.add_paragraph()
        p = doc.add_paragraph()
        amt = float(data.get('amount',0))
        r = p.add_run(f"Итого: {amt:,.2f} руб. {data.get('vat_rate','0')}% НДС")
        r.bold = True
        doc.save(output_path)
        return True
    except Exception:
        return False


def _generate_invoice_python(data: dict, output_path: str) -> bool:
    return _generate_act_python(data, output_path)


def build_doc_data(db, trip_id=None, doc_number="1", doc_date=None,
                   counterparty_id=None, service_desc=None, amount=None,
                   vat_rate=0, vat_amount=0, amount_no_vat=None):
    """Build data dict from DB for document generation."""
    ci = db.get_active_company()
    data = {
        "act_number":        doc_number,
        "date":              doc_date or datetime.now().strftime("%d.%m.%Y"),
        "contractor_name":   ci.get("name",""),
        "contractor_inn":    ci.get("inn",""),
        "contractor_kpp":    ci.get("kpp",""),
        "contractor_ogrn":   ci.get("ogrn",""),
        "contractor_address": ci.get("address",""),
        "contractor_address2": ci.get("address_fact",""),
        "contractor_bank":   ci.get("bank",""),
        "contractor_rs":     ci.get("rs",""),
        "contractor_bik":    ci.get("bik",""),
        "contractor_ks":     ci.get("ks",""),
        "contractor_phone":  ci.get("phone",""),
        "contractor_email":  ci.get("email",""),
        "contractor_director": ci.get("director",""),
        "client_name":       "",
        "client_inn":        "",
        "client_kpp":        "",
        "client_address":    "",
        "service_desc":      service_desc or "",
        "amount":            str(amount or 0),
        "vat_rate":          str(vat_rate or 0),
        "vat_amount":        str(vat_amount or 0),
        "amount_no_vat":     str(amount_no_vat or amount or 0),
        "base_doc":          "",
    }

    # Fill from trip
    if trip_id:
        trip = db.get_by_id("trips", trip_id)
        if trip:
            data["service_desc"]  = trip.get("route", "")
            data["amount"]        = str(trip.get("income_total", 0))
            data["vat_rate"]      = str(trip.get("vat_rate", 0))
            data["vat_amount"]    = str(trip.get("vat_amount", 0))
            data["amount_no_vat"] = str(trip.get("income_no_vat", 0))
            data["date"]          = trip.get("date", data["date"])
            data["act_number"]    = trip.get("act_number", doc_number)
            data["base_doc"]      = f"по заявке от {trip.get('date','')}"
            if trip.get("counterparty_id"):
                counterparty_id = trip["counterparty_id"]

    # Fill counterparty
    if counterparty_id:
        cp = db.get_by_id("counterparties", counterparty_id)
        if cp:
            data["client_name"]    = cp.get("name","")
            data["client_inn"]     = cp.get("inn","")
            data["client_kpp"]     = cp.get("kpp","")
            data["client_address"] = cp.get("address","")
            data["client_bank"]    = cp.get("bank","")
            data["client_rs"]      = cp.get("rs","")
            data["client_bik"]     = cp.get("bik","")
            data["client_ks"]      = cp.get("ks","")

    return data
