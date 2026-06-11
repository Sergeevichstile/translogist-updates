import tkinter as tk
from tkinter import messagebox, filedialog
import os
from datetime import datetime
from .ui_helpers import *

DOCX_OK = None  # Lazy check

def _check_docx():
    global DOCX_OK
    if DOCX_OK is None:
        try:
            from docx import Document
            DOCX_OK = True
        except ImportError:
            DOCX_OK = False
    return DOCX_OK

TEMPLATES = {
    "Договор с водителем":      "contract_driver",
    "Договор с контрагентом":   "contract_client",
    "Путевой лист":             "waybill",
    "Акт выполненных работ":    "act",
    "Счёт-фактура / УПД":       "invoice",
    "Заказ-наряд":              "work_order",
    "Счёт на оплату":           "payment_invoice",
}


class DocumentsFrame(tk.Frame):
    def __init__(self, parent, db):
        super().__init__(parent, bg=BG)
        self.db = db
        self._build()

    def _build(self):
        hdr = tk.Frame(self, bg=BG, padx=24, pady=14)
        hdr.pack(fill="x")
        label(hdr, "📄 Документы", size=16, bold=True, bg=BG).pack(side="left")
        self.active_co_label = label(hdr, "", bg=BG, color=SUCCESS, size=10)
        self.active_co_label.pack(side="right", padx=(0,12))
        btn(hdr, "⚙️  Компании", self._open_companies, color=CARD).pack(side="right")

        if not _check_docx():
            w = tk.Frame(self, bg="#7c2d12", padx=20, pady=10)
            w.pack(fill="x", padx=24)
            label(w, "⚠  pip install python-docx", bg="#7c2d12", color="#fef2f2").pack()

        body = tk.Frame(self, bg=BG, padx=24)
        body.pack(fill="both", expand=True, pady=8)

        # Left — template list
        left = tk.Frame(body, bg=CARD, width=210, padx=12, pady=12)
        left.pack(side="left", fill="y", padx=(0,16))
        left.pack_propagate(False)
        label(left, "Тип документа", bold=True, bg=CARD).pack(anchor="w", pady=(0,10))
        self.tmpl_var = tk.StringVar(value=list(TEMPLATES.keys())[0])
        for name in TEMPLATES:
            tk.Radiobutton(left, text=name, variable=self.tmpl_var, value=name,
                           bg=CARD, fg=TEXT, selectcolor=BG2,
                           activebackground=CARD, font=("Arial",10),
                           command=self._on_template_change).pack(anchor="w", pady=3)

        # Recent docs
        label(left, "Последние документы", bold=True, bg=CARD, size=10, color=MUTED).pack(
            anchor="w", pady=(16,6))
        self.recent_frame = tk.Frame(left, bg=CARD)
        self.recent_frame.pack(fill="x")

        # Right panel
        self.right_outer = tk.Frame(body, bg=BG)
        self.right_outer.pack(side="left", fill="both", expand=True)

        # Fixed generate button at bottom - always visible
        self.btn_frame = tk.Frame(self.right_outer, bg="#1e2435", padx=16, pady=10)
        self.btn_frame.pack(fill="x", side="bottom")
        btn(self.btn_frame, "📄  Создать документ (.docx)", self._generate,
            color=ACCENT).pack(side="left", ipadx=10)

        # Scrollable form area above button
        scroll_area = tk.Frame(self.right_outer, bg=BG)
        scroll_area.pack(fill="both", expand=True)

        self.canvas = tk.Canvas(scroll_area, bg=BG, highlightthickness=0)
        sb = ttk.Scrollbar(scroll_area, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        self.right = tk.Frame(self.canvas, bg=BG)
        self.canvas_window = self.canvas.create_window((0,0), window=self.right, anchor="nw")
        self.right.bind("<Configure>", lambda e: self.canvas.configure(
            scrollregion=self.canvas.bbox("all")))
        self.canvas.bind("<Configure>", lambda e: self.canvas.itemconfig(
            self.canvas_window, width=e.width))
        self.canvas.bind("<MouseWheel>", lambda e: self.canvas.yview_scroll(
            int(-1*(e.delta/120)), "units"))

        self._build_form()

    def _on_trip_select(self, e=None):
        """Fill form fields from selected trip."""
        try:
            idx = list(self.trip_var.get())
            # Find selected trip
            trips = self.db.get_all("trips")
            cp_map = {c["id"]: c.get("name","") for c in self.db.get_all("counterparties")}
            trip_labels = [
                f"{t.get('date','')} | {t.get('route','')[:40]} | {cp_map.get(t.get('counterparty_id',''),'')}"
                for t in trips
            ]
            sel = self.trip_var.get()
            trip = None
            for i, lbl in enumerate(trip_labels):
                if sel == lbl:
                    trip = trips[i]
                    break
            if not trip: return
            # Fill route
            if "route" in self.form_vars:
                self.form_vars["route"].set(trip.get("route",""))
            # Fill date
            if "doc_date" in self.form_vars:
                self.form_vars["doc_date"].set(trip.get("date",""))
            # Fill act number
            if "doc_number" in self.form_vars and trip.get("act_number"):
                self.form_vars["doc_number"].set(trip.get("act_number",""))
            # Fill financial fields
            if "amount_no_vat" in self.form_vars:
                self.form_vars["amount_no_vat"].set(str(trip.get("income_no_vat","")))
            if hasattr(self, "vat_mode"):
                self.vat_mode.set(str(int(trip.get("vat_rate",0))))
            # Fill counterparty
            if "counterparty" in self.form_vars and trip.get("counterparty_id"):
                cp_name = cp_map.get(trip["counterparty_id"],"")
                if cp_name: self.form_vars["counterparty"].set(cp_name)
        except Exception as ex:
            pass

    def _on_template_change(self):
        for w in self.right.winfo_children():
            w.destroy()
        self._build_form()

    def _open_companies(self):
        import subprocess, sys, os
        # Just show info
        from tkinter import messagebox
        messagebox.showinfo("Реквизиты", "Перейдите в раздел '🏦 Компании' для управления реквизитами")

    def _build_form(self):
        from tkinter import ttk as _ttk
        self.ttk = _ttk
        tmpl = self.tmpl_var.get()
        label(self.right, f"📝  {tmpl}", size=13, bold=True, bg=BG).pack(anchor="w", pady=(0,12))

        self.form_vars = {}
        self.vat_mode = tk.StringVar(value="5")

        self._field_row("Дата документа:", "doc_date", datetime.now().strftime("%d.%m.%Y"))
        self._field_row("Номер документа:", "doc_number", "")

        drivers   = self.db.get_all("drivers")
        trucks    = self.db.get_all("trucks")
        cps       = self.db.get_all("counterparties")
        carriers  = self.db.get_all("carriers")

        if tmpl in ["Договор с водителем","Путевой лист","Акт выполненных работ","Заказ-наряд"]:
            self._select_row("Водитель:", "driver",
                             [d.get("name","") for d in drivers], drivers)

        if tmpl in ["Путевой лист","Акт выполненных работ"]:
            self._select_row("Фура:", "truck",
                             [t.get("plate","") for t in trucks], trucks)

        if tmpl in ["Договор с контрагентом","Акт выполненных работ",
                    "Счёт-фактура / УПД","Заказ-наряд","Счёт на оплату","Договор с водителем"]:
            self._select_row("Контрагент:", "counterparty",
                             [c.get("name","") for c in cps], cps)

        if tmpl in ["Путевой лист","Акт выполненных работ","Счёт-фактура / УПД","Заказ-наряд"]:
            # Trip selector for auto-fill
            trips = self.db.get_all("trips")
            if trips:
                tk.Frame(self.right, bg=BG, height=4).pack()
                trip_card = tk.Frame(self.right, bg="#0f2027", padx=10, pady=8)
                trip_card.pack(fill="x", pady=(0,8))
                label(trip_card, "📋  Заполнить из рейса (необязательно):",
                      bg="#0f2027", color="#7dd3fc", size=10).pack(anchor="w", pady=(0,4))
                trip_names = ["— Выбрать рейс —"] + [
                    f"{t.get('date','')} | {t.get('act_number','—')} | {t.get('route','')[:40]}"
                    for t in trips
                ]
                self.trip_ids = [None] + [t["id"] for t in trips]
                self.trip_var = tk.StringVar()
                trip_cb = combo(trip_card, trip_names, textvariable=self.trip_var, width=52)
                trip_cb.current(0)
                trip_cb.pack(anchor="w")
                trip_cb.bind("<<ComboboxSelected>>", self._on_trip_select)

        if tmpl == "Путевой лист":
            self._field_row("Маршрут:", "route", "")
            self._field_row("Одометр (выезд, км):", "odo_start", "")
            self._field_row("Одометр (возврат, км):", "odo_end", "")
            self._field_row("Топливо выдано (л):", "fuel_given", "")
            self._field_row("Топливо сдано (л):", "fuel_returned", "")

        if tmpl in ["Акт выполненных работ","Счёт-фактура / УПД","Счёт на оплату"]:
            self._field_row("Описание услуг:", "service_desc", "Услуги по перевозке грузов")
            self._field_row("Сумма без НДС (₽):", "amount_no_vat", "")
            # VAT selector
            vf_label = tk.Frame(self.right, bg=BG)
            vf_label.pack(anchor="w", pady=(4,0))
            label(vf_label, "НДС:", bg=BG, size=10).pack(side="left", width=200, anchor="w")
            vf = tk.Frame(self.right, bg=BG)
            vf.pack(anchor="w", pady=2)
            for v, t in [("0","Без НДС"),("5","НДС 5%")]:
                tk.Radiobutton(vf, text=t, variable=self.vat_mode, value=v,
                               bg=BG, fg=TEXT, selectcolor=CARD, activebackground=BG,
                               command=self._recalc_vat).pack(side="left", padx=6)
            self.vat_result_label = label(self.right, "", bg=BG, color=WARN, size=10)
            self.vat_result_label.pack(anchor="w")

        if tmpl == "Заказ-наряд":
            self._field_row("Вид работ:", "work_type", "")
            self._field_row("Стоимость работ (₽):", "work_cost", "")

        if tmpl in ["Договор с водителем","Договор с контрагентом"]:
            self._field_row("Срок договора (дней):", "contract_days", "365")
            self._field_row("Предмет договора:", "contract_subject",
                            "Оказание транспортных услуг по перевозке грузов")

        tk.Frame(self.right, bg=BG, height=20).pack()

    def _field_row(self, lbl_text, key, default=""):
        row = tk.Frame(self.right, bg=BG)
        row.pack(fill="x", pady=3)
        lbl_w = label(row, lbl_text, bg=BG, size=10)
        lbl_w.pack(side="left", anchor="w")
        lbl_w.configure(width=28)
        var = tk.StringVar(value=default)
        e = entry(row, width=30, textvariable=var)
        e.pack(side="left")
        self.form_vars[key] = var

    def _select_row(self, lbl_text, key, names, records):
        row = tk.Frame(self.right, bg=BG)
        row.pack(fill="x", pady=3)
        lbl_w = label(row, lbl_text, bg=BG, size=10)
        lbl_w.pack(side="left", anchor="w")
        lbl_w.configure(width=28)
        var = tk.StringVar()
        cb = combo(row, names or ["—"], textvariable=var, width=30)
        if names: cb.current(0)
        cb.pack(side="left")
        self.form_vars[key] = var
        setattr(self, f"_{key}_records", records)

    def _on_trip_select(self, event=None):
        """Auto-fill fields from selected trip."""
        if not hasattr(self, "trip_ids") or not hasattr(self, "trip_var"):
            return
        sel = self.trip_var.get()
        idx = None
        trips = self.db.get_all("trips")
        trip_names = ["— Выбрать рейс —"] + [
            f"{t.get('date','')} | {t.get('act_number','—')} | {t.get('route','')[:40]}"
            for t in trips
        ]
        if sel in trip_names:
            i = trip_names.index(sel)
            if i > 0:
                trip = trips[i-1]
                # Fill route
                if "route" in self.form_vars:
                    self.form_vars["route"].set(trip.get("route",""))
                # Fill doc number from act
                if "doc_number" in self.form_vars and trip.get("act_number"):
                    self.form_vars["doc_number"].set(trip.get("act_number",""))
                # Fill date
                if "doc_date" in self.form_vars and trip.get("date"):
                    self.form_vars["doc_date"].set(trip.get("date",""))
                # Fill amount
                if "amount_no_vat" in self.form_vars:
                    self.form_vars["amount_no_vat"].set(str(trip.get("income_no_vat","")))
                if hasattr(self, "vat_mode"):
                    self.vat_mode.set(str(int(trip.get("vat_rate",0))))
                # Fill driver/truck from trip
                if trip.get("driver_id") and "driver" in self.form_vars:
                    drivers = self.db.get_all("drivers")
                    drv = next((d for d in drivers if d["id"]==trip["driver_id"]), None)
                    if drv: self.form_vars["driver"].set(drv.get("name",""))
                if trip.get("truck_id") and "truck" in self.form_vars:
                    trucks = self.db.get_all("trucks")
                    trk = next((t for t in trucks if t["id"]==trip["truck_id"]), None)
                    if trk: self.form_vars["truck"].set(trk.get("plate",""))
                self._recalc_vat()

    def _recalc_vat(self):
        try:
            amt = float(self.form_vars["amount_no_vat"].get().replace(",","."))
            rate = float(self.vat_mode.get())
            vat = round(amt * rate / 100, 2)
            total = round(amt + vat, 2)
            if rate > 0:
                self.vat_result_label.config(
                    text=f"  НДС {rate:.0f}%: {vat:,.2f} ₽   |   Итого с НДС: {total:,.2f} ₽")
            else:
                self.vat_result_label.config(text="  Без НДС")
        except Exception:
            if hasattr(self, "vat_result_label"):
                self.vat_result_label.config(text="")

    def _get_val(self, key, default=""):
        v = self.form_vars.get(key)
        return v.get().strip() if v else default

    def _generate(self):
        tmpl    = self.tmpl_var.get()
        doc_num = self._get_val("doc_number","1")
        doc_date= self._get_val("doc_date")
        fname   = f"{TEMPLATES[tmpl]}_{doc_num}_{doc_date.replace('.','_')}.docx"

        out_dir = "output"
        os.makedirs(out_dir, exist_ok=True)

        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word документ","*.docx")],
            initialfile=fname, initialdir=out_dir)
        if not save_path: return

        # Try new high-quality generators first
        from .doc_generator import generate_act, generate_invoice, generate_upd, build_doc_data

        # Get selected trip id
        trip_id = None
        if hasattr(self, 'trip_ids') and hasattr(self, 'trip_var'):
            try:
                trips = self.db.get_all("trips")
                cp_map = {c["id"]: c.get("name","") for c in self.db.get_all("counterparties")}
                sel = self.trip_var.get()
                for i, t in enumerate(trips):
                    lbl = f"{t.get('date','')} | {t.get('act_number','—')} | {t.get('route','')[:40]}"
                    if sel == lbl:
                        trip_id = t["id"]
                        break
            except Exception:
                pass

        # Get counterparty
        cp_id = None
        cp_name = self._get_val("counterparty","")
        if cp_name:
            for c in self.db.get_all("counterparties"):
                if c.get("name") == cp_name:
                    cp_id = c["id"]; break

        # Build financial data
        try:
            amt = float(self._get_val("amount_no_vat","0").replace(",",""))
            vat_r = float(self.vat_mode.get()) if hasattr(self,"vat_mode") else 0
            vat_a = round(amt * vat_r / 100, 2)
            total = round(amt + vat_a, 2)
        except Exception:
            amt = vat_r = vat_a = total = 0

        svc = self._get_val("service_desc","Услуги по перевозке грузов")

        data = build_doc_data(
            self.db,
            trip_id=trip_id,
            doc_number=doc_num,
            doc_date=doc_date,
            counterparty_id=cp_id,
            service_desc=svc,
            amount=total or amt,
            vat_rate=vat_r,
            vat_amount=vat_a,
            amount_no_vat=amt,
        )

        ok = False
        if tmpl in ["Акт выполненных работ"]:
            ok = generate_act(data, save_path)
        elif tmpl in ["Счёт на оплату"]:
            ok = generate_invoice(data, save_path)
        elif tmpl in ["Счёт-фактура / УПД"]:
            ok = generate_upd(data, save_path)
        else:
            # Fallback to python-docx for other templates
            if not DOCX_OK:
                messagebox.showerror("Ошибка","Установите python-docx"); return
            subs = self._build_subs()
            doc  = self._make_doc(tmpl, subs)
            if doc:
                doc.save(save_path)
                ok = True

        if ok:
            self.db.add("documents", {
                "type": tmpl, "number": doc_num,
                "date": doc_date, "path": save_path
            })
            self.refresh()
            messagebox.showinfo("Готово", f"Документ сохранён:\n{save_path}")
        else:
            messagebox.showerror("Ошибка", "Не удалось создать документ.\nУбедитесь что Node.js установлен или используйте другой тип документа.")

    def _build_subs(self):
        subs = {
            "{{ДАТА}}":         self._get_val("doc_date"),
            "{{НОМЕР}}":        self._get_val("doc_number","1"),
            "{{МАРШРУТ}}":      self._get_val("route"),
            "{{ОДОМЕТР_СТАРТ}}": self._get_val("odo_start"),
            "{{ОДОМЕТР_КОНЕЦ}}": self._get_val("odo_end"),
            "{{ТОПЛИВО_ВЫДАНО}}": self._get_val("fuel_given"),
            "{{ТОПЛИВО_СДАНО}}": self._get_val("fuel_returned"),
            "{{ОПИСАНИЕ_УСЛУГ}}": self._get_val("service_desc"),
            "{{ВИД_РАБОТ}}":    self._get_val("work_type"),
            "{{СТОИМОСТЬ_РАБОТ}}": self._get_val("work_cost"),
            "{{СРОК_ДОГОВОРА}}": self._get_val("contract_days","365"),
            "{{ПРЕДМЕТ_ДОГОВОРА}}": self._get_val("contract_subject"),
        }
        # Company info — from active company
        subs.update(self.db.get_active_company_subs())
        # Driver
        drv_name = self._get_val("driver")
        if drv_name:
            drivers = self.db.get_all("drivers")
            d = next((x for x in drivers if x.get("name")==drv_name), {})
            subs.update({
                "{{ФИО_ВОДИТЕЛЯ}}":   d.get("name", drv_name),
                "{{ПАСПОРТ_ВОДИТЕЛЯ}}": d.get("passport",""),
                "{{ТЕЛЕФОН_ВОДИТЕЛЯ}}": d.get("phone",""),
                "{{ВУ_ВОДИТЕЛЯ}}":    d.get("license",""),
            })
        # Truck
        plate = self._get_val("truck")
        if plate:
            trucks = self.db.get_all("trucks")
            t = next((x for x in trucks if x.get("plate")==plate), {})
            subs.update({
                "{{ГОС_НОМЕР}}":  t.get("plate", plate),
                "{{МАРКА_ТС}}":   t.get("model",""),
                "{{ГОД_ТС}}":     str(t.get("year","")),
            })
        # Counterparty
        cp_name = self._get_val("counterparty")
        if cp_name:
            cps = self.db.get_all("counterparties")
            c = next((x for x in cps if x.get("name")==cp_name), {})
            subs.update({
                "{{КОНТРАГЕНТ}}":     c.get("name", cp_name),
                "{{ИНН_КОНТРАГЕНТА}}": c.get("inn",""),
                "{{КПП_КОНТРАГЕНТА}}": c.get("kpp",""),
                "{{АДРЕС_КОНТРАГЕНТА}}": c.get("address",""),
                "{{КОНТАКТ_КОНТРАГЕНТА}}": c.get("contact",""),
            })
        # VAT
        try:
            amt = float(self._get_val("amount_no_vat","0").replace(",","."))
            rate = float(self.vat_mode.get())
        except Exception:
            amt = rate = 0
        vat_amt = round(amt * rate / 100, 2)
        total   = round(amt + vat_amt, 2)
        subs.update({
            "{{СУММА_БЕЗ_НДС}}": f"{amt:,.2f}",
            "{{НДС_ПРОЦЕНТ}}":   f"{rate:.0f}",
            "{{НДС_СУММА}}":     f"{vat_amt:,.2f}",
            "{{СУММА_ИТОГО}}":   f"{total:,.2f}",
        })
        return subs

    def _make_doc(self, tmpl_name, subs):
        tmpl_key  = TEMPLATES[tmpl_name]
        tmpl_path = os.path.join("templates", f"{tmpl_key}.docx")
        if os.path.exists(tmpl_path) and _check_docx():
            from docx import Document as _Doc
            doc = _Doc(tmpl_path)
            _fill_doc(doc, subs)
            return doc
        return _create_doc(tmpl_name, subs)

    def refresh(self):
        active = self.db.get_active_company()
        co_name = active.get("name","Не выбрана") if active else "Не выбрана"
        if hasattr(self, "active_co_label"):
            self.active_co_label.config(text=f"Активная: {co_name}")
        for w in self.recent_frame.winfo_children():
            w.destroy()
        docs = list(reversed(self.db.get_all("documents")))[:5]
        for d in docs:
            f = tk.Frame(self.recent_frame, bg=CARD)
            f.pack(fill="x", pady=2)
            label(f, f"№{d.get('number','')} {d.get('type','')}",
                  bg=CARD, size=9, color=MUTED).pack(anchor="w")


# ── Company dialog ─────────────────────────────────────────────────────────

class CompanyDialog(tk.Toplevel):
    def __init__(self, parent, db):
        super().__init__(parent)
        self.db = db
        self.title("Реквизиты компании")
        self.geometry("460x500")
        self.configure(bg=BG)
        self.resizable(False, False)
        self.grab_set()
        self._build()

    def _build(self):
        ci = self.db.get_company_info()
        fields = [
            ("Название организации:", "name"),
            ("ИНН:", "inn"),
            ("КПП:", "kpp"),
            ("ОГРН:", "ogrn"),
            ("Юридический адрес:", "address"),
            ("Банк:", "bank"),
            ("Расчётный счёт:", "rs"),
            ("БИК:", "bik"),
            ("Директор (ФИО):", "director"),
            ("Телефон:", "phone"),
        ]
        self.entries = {}
        for i, (lbl_text, key) in enumerate(fields):
            label(self, lbl_text, bg=BG, size=10).grid(
                row=i, column=0, sticky="w", padx=20, pady=5)
            e = entry(self, width=28)
            e.insert(0, ci.get(key,""))
            e.grid(row=i, column=1, padx=20, pady=5)
            self.entries[key] = e
        btn(self, "💾 Сохранить", self._save, color=SUCCESS).grid(
            row=len(fields), column=0, columnspan=2, pady=16)

    def _save(self):
        data = {k: e.get().strip() for k, e in self.entries.items()}
        self.db.set_company_info(data)
        messagebox.showinfo("Готово","Реквизиты сохранены!")
        self.destroy()


# ── Doc helpers ────────────────────────────────────────────────────────────

def _fill_doc(doc, subs):
    def replace_in_para(para):
        full = para.text
        for k, v in subs.items():
            full = full.replace(k, v)
        if full != para.text:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = full
            else:
                para.add_run(full)
    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)


def _h(doc, text, size=13, bold=True, center=False):
    p = doc.add_paragraph()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    return p


def _line(doc, label_text, value, size=11):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label_text}: "); r1.bold = True; r1.font.size = Pt(size)
    r2 = p.add_run(value or "________________"); r2.font.size = Pt(size)


def g(subs, key): return subs.get(key, "")


def _create_doc(tmpl_name, subs):
    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(1.5)

    _h(doc, tmpl_name.upper(), size=14, center=True)
    _h(doc, f"№ {g(subs,'{{НОМЕР}}')}  от  {g(subs,'{{ДАТА}}')}",
       size=11, bold=False, center=True)
    doc.add_paragraph()

    # Company block
    ci_name = g(subs, "{{КОМПАНИЯ}}")
    if ci_name and ci_name != "________________":
        _h(doc, "ИСПОЛНИТЕЛЬ", size=11)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.cell(0,0).text = (
            f"Организация: {g(subs,'{{КОМПАНИЯ}}')}\n"
            f"ИНН: {g(subs,'{{ИНН_КОМПАНИИ}}')}  КПП: {g(subs,'{{КПП_КОМПАНИИ}}')}\n"
            f"ОГРН: {g(subs,'{{ОГРН_КОМПАНИИ}}')}\n"
            f"Адрес: {g(subs,'{{АДРЕС_КОМПАНИИ}}')}"
        )
        tbl.cell(0,1).text = (
            f"Банк: {g(subs,'{{БАНК_КОМПАНИИ}}')}\n"
            f"Р/с: {g(subs,'{{РС_КОМПАНИИ}}')}\n"
            f"БИК: {g(subs,'{{БИК_КОМПАНИИ}}')}"
        )
        doc.add_paragraph()

    # Counterparty
    if g(subs,"{{КОНТРАГЕНТ}}"):
        _h(doc, "ЗАКАЗЧИК", size=11)
        _line(doc, "Организация", g(subs,"{{КОНТРАГЕНТ}}"))
        _line(doc, "ИНН", g(subs,"{{ИНН_КОНТРАГЕНТА}}"))
        _line(doc, "Адрес", g(subs,"{{АДРЕС_КОНТРАГЕНТА}}"))
        doc.add_paragraph()

    # Driver/Truck
    if g(subs,"{{ФИО_ВОДИТЕЛЯ}}"):
        _h(doc, "ВОДИТЕЛЬ", size=11)
        _line(doc, "ФИО", g(subs,"{{ФИО_ВОДИТЕЛЯ}}"))
        _line(doc, "Паспорт", g(subs,"{{ПАСПОРТ_ВОДИТЕЛЯ}}"))
        _line(doc, "Вод. удостоверение", g(subs,"{{ВУ_ВОДИТЕЛЯ}}"))
        _line(doc, "Телефон", g(subs,"{{ТЕЛЕФОН_ВОДИТЕЛЯ}}"))
        doc.add_paragraph()

    if g(subs,"{{ГОС_НОМЕР}}"):
        _h(doc, "ТРАНСПОРТНОЕ СРЕДСТВО", size=11)
        _line(doc, "Марка/модель", g(subs,"{{МАРКА_ТС}}"))
        _line(doc, "Гос. номер", g(subs,"{{ГОС_НОМЕР}}"))
        if g(subs,"{{МАРШРУТ}}"): _line(doc,"Маршрут", g(subs,"{{МАРШРУТ}}"))
        if g(subs,"{{ОДОМЕТР_СТАРТ}}"): _line(doc,"Одометр выезд (км)", g(subs,"{{ОДОМЕТР_СТАРТ}}"))
        if g(subs,"{{ОДОМЕТР_КОНЕЦ}}"): _line(doc,"Одометр возврат (км)", g(subs,"{{ОДОМЕТР_КОНЕЦ}}"))
        if g(subs,"{{ТОПЛИВО_ВЫДАНО}}"): _line(doc,"Топливо выдано (л)", g(subs,"{{ТОПЛИВО_ВЫДАНО}}"))
        if g(subs,"{{ТОПЛИВО_СДАНО}}"): _line(doc,"Топливо сдано (л)", g(subs,"{{ТОПЛИВО_СДАНО}}"))
        doc.add_paragraph()

    # Financial
    amt_str = g(subs,"{{СУММА_БЕЗ_НДС}}")
    if amt_str and amt_str != "0,00":
        _h(doc, "СТОИМОСТЬ УСЛУГ", size=11)
        if g(subs,"{{ОПИСАНИЕ_УСЛУГ}}"): _line(doc,"Услуга", g(subs,"{{ОПИСАНИЕ_УСЛУГ}}"))
        if g(subs,"{{ВИД_РАБОТ}}"): _line(doc,"Вид работ", g(subs,"{{ВИД_РАБОТ}}"))

        # Financial table
        tbl = doc.add_table(rows=4, cols=2)
        tbl.style = "Table Grid"
        rows_data = [
            ("Сумма без НДС", f"{g(subs,'{{СУММА_БЕЗ_НДС}}')} ₽"),
            (f"НДС {g(subs,'{{НДС_ПРОЦЕНТ}}')}%", f"{g(subs,'{{НДС_СУММА}}')} ₽"),
            ("ИТОГО к оплате", f"{g(subs,'{{СУММА_ИТОГО}}')} ₽"),
            ("Оплата", "□ Наличные   □ Безналичные"),
        ]
        for i,(k,v) in enumerate(rows_data):
            tbl.cell(i,0).text = k
            tbl.cell(i,1).text = v
            tbl.cell(i,0).paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

    if g(subs,"{{ПРЕДМЕТ_ДОГОВОРА}}"):
        _h(doc,"ПРЕДМЕТ ДОГОВОРА", size=11)
        p = doc.add_paragraph()
        p.add_run(g(subs,"{{ПРЕДМЕТ_ДОГОВОРА}}")).font.size = Pt(11)
        if g(subs,"{{СРОК_ДОГОВОРА}}"): _line(doc,"Срок договора (дней)", g(subs,"{{СРОК_ДОГОВОРА}}"))
        doc.add_paragraph()

    # Signatures
    doc.add_paragraph()
    _h(doc,"ПОДПИСИ СТОРОН", size=11)
    sig_tbl = doc.add_table(rows=3, cols=2)
    sig_tbl.style = "Table Grid"
    sig_tbl.cell(0,0).text = f"ИСПОЛНИТЕЛЬ: {g(subs,'{{КОМПАНИЯ}}')}"
    sig_tbl.cell(0,1).text = f"ЗАКАЗЧИК: {g(subs,'{{КОНТРАГЕНТ}}')}"
    sig_tbl.cell(1,0).text = f"Директор: {g(subs,'{{ДИРЕКТОР}}')}"
    sig_tbl.cell(1,1).text = "Директор: ________________"
    sig_tbl.cell(2,0).text = "Подпись: ________________  М.П."
    sig_tbl.cell(2,1).text = "Подпись: ________________  М.П."

    return doc
